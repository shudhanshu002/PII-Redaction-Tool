import streamlit as st
import io
import re
import docx
from faker import Faker
import spacy

# ---------------------------------------------------------
# Application Configuration & UI Setup
# ---------------------------------------------------------
st.set_page_config(page_title="PII Redaction Engine", page_icon="🛡️", layout="centered")

st.title("🛡️ PII Redaction Tool for Word Documents")
st.markdown("""
Upload a Microsoft Word (`.docx`) document to sanitize sensitive personal details, contact information, financial identifiers, and embedded images.
""")

# Cache SpaCy model loading across sessions
@st.cache_resource
def load_nlp_model():
    return spacy.load("en_core_web_sm")

nlp = load_nlp_model()
fake = Faker("en_IN")

# ---------------------------------------------------------
# Document Redactor Engine
# ---------------------------------------------------------
class DocumentPIIRedactor:
    def __init__(self):
        self.mapping = {}

        # Gazetteer of domain entities (promoters, trusts, banks, registrars)
        self.known_entities = {
            "KSH INTERNATIONAL LIMITED", "KSH International Limited", "KSH INTERNATIONAL PRIVATE LIMITED",
            "Bhandary Metal Extrusion Private Limited", "KUSHAL SUBBAYYA HEGDE", "Kushal Subbayya Hegde",
            "PUSHPA KUSHAL HEGDE", "Pushpa Kushal Hegde", "RAJESH KUSHAL HEGDE", "Rajesh Kushal Hegde",
            "ROHIT KUSHAL HEGDE", "Rohit Kushal Hegde", "RAKHI GIRIJA SHETTY", "Rakhi Girija Shetty",
            "Sarthak Malvadkar", "Sandesh Bhagwat", "Amod Joshi", "Sangeeta Ramprasad Rai",
            "Maithili Rajesh Hegde", "Katyayani Balasubramanian", "Lalit Muljibhai Sarvaiya",
            "Dinesh Hirachand Munot", "Ajay Shriram Patil", "Ram Kumar Tiwari", "Indu Jacob",
            "DHAULAGIRI FAMILY TRUST", "EVEREST FAMILY TRUST", "MAKALU FAMILY TRUST", 
            "BROAD FAMILY TRUST", "ANNAPURNA FAMILY TRUST", "KANCHENJUNGA FAMILY TRUST",
            "WATERLOO INDUSTRIAL PARK VI PRIVATE LIMITED", "Dhaulagiri Family Trust",
            "Everest Family Trust", "Makalu Family Trust", "Broad Family Trust",
            "Annapurna Family Trust", "Kanchenjunga Family Trust",
            "ICICI Securities Limited", "ICICI Bank", "HDFC Bank Limited", "HDFC Bank",
            "Link Intime India Private Limited", "Nuvama", "State Bank of India",
            "Bajaj Finance", "Federal Bank", "Export Import Bank of India", "Citibank",
            "ICICI", "HDFC", "Trilegal"
        }

        # Regular expressions for structured entities
        self.patterns = {
            "URL": r'(?i)(?:https?://|www\.)[a-z0-9.\-]+(?:\s*[\./]\s*[a-z0-9\-/]+)*',
            "EMAIL": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            "PHONE": r'(?:\+\s*)?(?:91[\s-]?)?(?:\(?\d{2,4}\)?[\s-]?)?\d{3,5}[\s-]?\d{4,5}\b',
            "PAN_CARD": r'\b[A-Z]{5}[0-9]{4}[A-Z]\b',
            "AADHAAR": r'\b[2-9]\d{3}\s?\d{4}\s?\d{4}\b',
        }

        # Protected terms to avoid false positives
        self.protected_terms = {
            "equity shares", "equity share", "offer for sale", "fresh issue", "promoter selling shareholders",
            "sebi icdr regulations", "book building process", "red herring prospectus", "prospectus",
            "company", "issuer", "offer price", "price band", "bid/offer period", "general risks",
            "risks in relation to the first offer", "stock exchanges", "bse limited", "national stock exchange",
            "qualified institutional buyers", "non-institutional investors", "retail individual investors",
            "companies act", "scrr", "asba", "upi", "roc", "drhp", "rhp", "bse", "nse", "sebi", "offer", "board",
            "bidders", "bidder", "statutory auditors", "statutory auditor", "chartered accountants"
        }

    def get_fake_value(self, pii_type: str, original_text: str) -> str:
        key = original_text.strip().lower()
        if key in self.mapping:
            return self.mapping[key]

        if pii_type == "URL":
            val = fake.url()
        elif pii_type == "EMAIL":
            val = fake.unique.email()
        elif pii_type == "PHONE":
            val = f"+91 {fake.msisdn()[3:13]}"
        elif pii_type == "PAN_CARD":
            val = f"{fake.lexify('?????').upper()}{fake.numerify('####')}{fake.lexify('?').upper()}"
        elif pii_type == "AADHAAR":
            val = f"{fake.random_int(2000, 9999)} {fake.random_int(1000, 9999)} {fake.random_int(1000, 9999)}"
        elif pii_type in ["PERSON", "KNOWN_PERSON"]:
            val = fake.unique.name()
        elif pii_type in ["COMPANY", "KNOWN_COMPANY"]:
            val = fake.unique.company()
        else:
            val = "[REDACTED]"

        self.mapping[key] = val
        return val

    def redact_text(self, text: str) -> str:
        if not text.strip():
            return text

        # Step 1: Flexible gazetteer matching
        for entity in sorted(self.known_entities, key=len, reverse=True):
            flexible_pattern = r'\s+'.join(map(re.escape, entity.split()))
            pattern = re.compile(flexible_pattern, re.IGNORECASE)
            
            for match in pattern.findall(text):
                pii_type = "KNOWN_COMPANY" if any(w in entity.upper() for w in ["LIMITED", "PRIVATE", "TRUST", "PARK", "BANK", "FINANCE"]) else "KNOWN_PERSON"
                replacement = self.get_fake_value(pii_type, match)
                text = text.replace(match, replacement)

        # Step 2: Strict regex rules
        for pii_type, regex in self.patterns.items():
            pattern = re.compile(regex)
            matches = set(pattern.findall(text))
            for match in matches:
                match_str = match if isinstance(match, str) else "".join(match)
                if "PLC" in text or "U28129" in match_str:
                    continue
                if match_str.strip():
                    fake_val = self.get_fake_value(pii_type, match_str)
                    text = text.replace(match_str, fake_val)

        # Step 3: SpaCy Named Entity Recognition
        if not text.isupper():
            doc = nlp(text)
            spans = []
            for ent in doc.ents:
                clean_ent = ent.text.strip().lower()
                if any(protected in clean_ent for protected in self.protected_terms):
                    continue
                if ent.label_ == "PERSON":
                    spans.append((ent.start_char, ent.end_char, "PERSON", ent.text))
                elif ent.label_ == "ORG" and len(ent.text) > 3:
                    spans.append((ent.start_char, ent.end_char, "COMPANY", ent.text))

            spans.sort(key=lambda x: x[0], reverse=True)
            for start, end, pii_type, orig in spans:
                if orig.strip().lower() not in self.protected_terms:
                    replacement = self.get_fake_value(pii_type, orig)
                    text = text[:start] + replacement + text[end:]

        # Step 4: Formatting cleanup
        text = text.replace(" com", "")

        return text

    def remove_image_pii(self, doc):
        images_removed = 0
        blank_image = b'\x47\x49\x46\x38\x39\x61\x01\x00\x01\x00\x80\x00\x00\xff\xff\xff\x00\x00\x00!\xf9\x04\x01\x00\x00\x00\x00,\x00\x00\x00\x00\x01\x00\x01\x00\x00\x02\x02D\x01\x00;'
        
        for part in doc.part.package.parts:
            if part.content_type.startswith('image/'):
                part._blob = blank_image
                images_removed += 1
        return images_removed

    def process_docx_stream(self, input_stream):
        doc = docx.Document(input_stream)
        for p in doc.paragraphs:
            if p.text:
                p.text = self.redact_text(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text:
                            p.text = self.redact_text(p.text)
        
        images_wiped = self.remove_image_pii(doc)
        
        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        return output_stream, images_wiped


ProductionPIIRedactor = DocumentPIIRedactor

# ---------------------------------------------------------
# Streamlit Interface Actions
# ---------------------------------------------------------
uploaded_file = st.file_uploader("Upload Word Document (.docx)", type="docx")

if uploaded_file is not None:
    st.info("Document loaded successfully. Click below to start sanitization.")
    
    if st.button("Redact Document"):
        with st.spinner("Processing document text and replacing sensitive entities..."):
            redactor = DocumentPIIRedactor()
            
            # Process in-memory stream
            file_stream = io.BytesIO(uploaded_file.getvalue())
            processed_stream, images_wiped = redactor.process_docx_stream(file_stream)
            
            st.success(f"✅ Document successfully sanitized ({images_wiped} images cleared).")
            
            # Download button
            st.download_button(
                label="📥 Download Redacted Document",
                data=processed_stream,
                file_name=f"Redacted_{uploaded_file.name}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
