import re
import sys
import os
import docx
from faker import Faker
import spacy

# loading SpaCy English language model here
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("SpaCy model 'en_core_web_sm' not found locally. Downloading model...")
    from spacy.cli import download
    download("en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

fake = Faker("en_IN")


class DocumentPIIRedactor:
    def __init__(self):
        self.mapping = {}

        # gazetteer of known specific thingss
        self.known_entities = {
            "KSH INTERNATIONAL LIMITED", "KSH International Limited", "KSH INTERNATIONAL PRIVATE LIMITED",
            "Bhandary Metal Extrusion Private Limited", "KUSHAL SUBBAYYA HEGDE", "Kushal Subbayya Hegde",
            "PUSHPA KUSHAL HEGDE", "Pushpa Kushal Hegde", "RAJESH KUSHAL HEGDE", "Rajesh Kushal Hegde",
            "ROHIT KUSHAL HEGDE", "Rohit Kushal Hegde", "RAKHI GIRIJA SHETTY", "Rakhi Girija Shetty",
            "Sarthak Malvadkar", "Sandesh Bhagwat", "Amod Joshi", "Sangeeta Ramprasad Rai",
            "Maithili Rajesh Hegde", "Katyayani Balasubramanian", "Lalit Muljibhai Sarvaiya",
            "Dinesh Hirachand Munot", "Ajay Shriram Patil", "Ram Kumar Tiwari", "Indu Jacob",
            "DHAULAGIRI FAMILY TRUST", "EVEREST FAMILY TRUST", "MAKALU FAMILY TRUST", 
            "BROAD FAMILY TRUST", "ANNAPURNA FAMILY TRUST", "KANCHENJUNGA FAMILY TRUST",
            "WATERLOO INDUSTRIAL PARK VI PRIVATE LIMITED", "Dhaulagiri Family Trust",
            "Everest Family Trust", "Makalu Family Trust", "Broad Family Trust",
            "Annapurna Family Trust", "Kanchenjunga Family Trust",
            "ICICI Securities Limited", "ICICI Bank", "HDFC Bank Limited", "HDFC Bank",
            "Link Intime India Private Limited", "Nuvama", "State Bank of India",
            "Bajaj Finance", "Federal Bank", "Export Import Bank of India", "Citibank"
        }

        # Additional words
        self.known_entities.update({
            "ICICI", "HDFC", "Trilegal"
        })

        # Regexs
        self.patterns = {
            "URL": r'(?i)(?:https?://|www\.)[a-z0-9.\-]+(?:\s*[\./]\s*[a-z0-9\-/]+)*(?:\s*(?:com|in|co|org|net|gov|edu|io|info)\b)?',
            "EMAIL": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            "PHONE": r'(?:\+\s*)?(?:91[\s-]?)?(?:\(?\d{2,4}\)?[\s-]?)?\d{3,5}[\s-]?\d{4,5}\b',
            "PAN_CARD": r'\b[A-Z]{5}[0-9]{4}[A-Z]\b',
            "AADHAAR": r'\b[2-9]\d{3}\s?\d{4}\s?\d{4}\b',
            "IP_ADDRESS": r'\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b',
            "CREDIT_CARD": r'\b(?:\d[ -]*?){13,16}\b',
            "SSN": r'\b\d{3}-\d{2}-\d{4}\b',
            "DOB": r'\b(?:0[1-9]|1[0-2])[\/.-](?:0[1-9]|[12]\d|3[01])[\/.-](?:19|20)\d{2}\b|\b\d{4}[-/.](?:0[1-9]|1[02])[-/.](?:0[1-9]|[12]\d|[301])\b',
        }

        #Technical/financial terms to preserve
        self.protected_terms = {
            "equity shares", "equity share", "offer for sale", "fresh issue", "promoter selling shareholders",
            "sebi icdr regulations", "book building process", "red herring prospectus", "prospectus",
            "company", "issuer", "offer price", "price band", "bid/offer period", "general risks",
            "risks in relation to the first offer", "stock exchanges", "bse limited", "national stock exchange",
            "qualified institutional buyers", "non-institutional investors", "retail individual investors",
            "companies act", "scrr", "asba", "upi", "roc", "drhp", "rhp", "bse", "nse", "sebi", "offer", "board",
            "bidders", "bidder", "statutory auditors", "statutory auditor", "chartered accountants"
        }

    def get_fake_value(self, pii_type: str, original_text: str) -> str:
        """Returns a consistent replacement for the given text."""
        key = original_text.strip().lower()
        if key in self.mapping:
            return self.mapping[key]

        if pii_type == "URL":
            val = fake.url()
        elif pii_type == "EMAIL":
            val = fake.unique.email()
        elif pii_type == "PHONE":
            val = f"+91 {fake.msisdn()[3:13]}"
        elif pii_type == "PAN_CARD":
            val = f"{fake.lexify('?????').upper()}{fake.numerify('####')}{fake.lexify('?').upper()}"
        elif pii_type == "AADHAAR":
            val = f"{fake.random_int(2000, 9999)} {fake.random_int(1000, 9999)} {fake.random_int(1000, 9999)}"
        elif pii_type in ["PERSON", "KNOWN_PERSON"]:
            val = fake.unique.name()
        elif pii_type in ["COMPANY", "KNOWN_COMPANY"]:
            val = fake.unique.company()
        elif pii_type == "ADDRESS":
            val = fake.street_address().replace('\n', ', ')
        elif pii_type == "SSN":
            val = fake.ssn()
        elif pii_type == "CREDIT_CARD":
            val = fake.credit_card_number()
        elif pii_type == "DOB":
            val = fake.date_of_birth().strftime('%Y-%m-%d')
        elif pii_type == "IP_ADDRESS":
            val = fake.ipv4()
        else:
            val = "[REDACTED]"

        self.mapping[key] = val
        return val

    def redact_text(self, text: str) -> str:
        """Sanitizes text strings using gazetteer, regex, and SpaCy NER filtering."""
        if not text.strip():
            return text

        # Match known gazetteer entities--
        for entity in sorted(self.known_entities, key=len, reverse=True):
            flexible_pattern = r'\s+'.join(map(re.escape, entity.split()))
            pattern = re.compile(flexible_pattern, re.IGNORECASE)
            
            for match in pattern.findall(text):
                pii_type = "KNOWN_COMPANY" if any(w in entity.upper() for w in ["LIMITED", "PRIVATE", "TRUST", "PARK", "BANK", "FINANCE", "ICICI", "HDFC", "TRILEGAL"]) else "KNOWN_PERSON"
                replacement = self.get_fake_value(pii_type, match)
                text = text.replace(match, replacement)

        # match regex--
        for pii_type, regex in self.patterns.items():
            pattern = re.compile(regex)
            matches = set(pattern.findall(text))
            for match in matches:
                match_str = match if isinstance(match, str) else (match[0] if isinstance(match, tuple) else "".join(match))
                if "PLC" in text or "U28129" in match_str:
                    continue
                if match_str.strip():
                    fake_val = self.get_fake_value(pii_type, match_str)
                    text = text.replace(match_str, fake_val)

        text = re.sub(r'(\.[a-z]{2,4}/?)\s+com\b', r'\1', text, flags=re.IGNORECASE)

        if not text.isupper():
            doc = nlp(text)
            spans = []
            for ent in doc.ents:
                clean_ent = ent.text.strip().lower()
                if any(protected in clean_ent for protected in self.protected_terms):
                    continue

                if ent.label_ == "PERSON":
                    spans.append((ent.start_char, ent.end_char, "PERSON", ent.text))
                elif ent.label_ == "ORG" and len(ent.text) > 3:
                    spans.append((ent.start_char, ent.end_char, "COMPANY", ent.text))

            spans.sort(key=lambda x: x[0], reverse=True)
            for start, end, pii_type, orig in spans:
                if orig.strip().lower() not in self.protected_terms:
                    replacement = self.get_fake_value(pii_type, orig)
                    text = text[:start] + replacement + text[end:]

        # formatting
        text = text.replace(" com", "")

        return text

    def remove_image_pii(self, doc):
        """Replaces embedded images with a transparent 1x1 GIF binary."""
        images_removed = 0
        blank_image = b'\x47\x49\x46\x38\x39\x61\x01\x00\x01\x00\x80\x00\x00\xff\xff\xff\x00\x00\x00!\xf9\x04\x01\x00\x00\x00\x00,\x00\x00\x00\x00\x01\x00\x01\x00\x00\x02\x02D\x01\x00;'
        
        for part in doc.part.package.parts:
            if part.content_type.startswith('image/'):
                part._blob = blank_image
                images_removed += 1
                
        return images_removed

    def process_docx(self, input_path: str, output_path: str):
        """Processes a DOCX file: paragraph text, tables, headers, footers, and images."""
        if not os.path.exists(input_path):
            print(f"Error: Input file '{input_path}' not found.")
            return False

        doc = docx.Document(input_path)

        # Process headers & footers
        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    if p.text:
                        p.text = self.redact_text(p.text)
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text:
                                    p.text = self.redact_text(p.text)
            if section.footer:
                for p in section.footer.paragraphs:
                    if p.text:
                        p.text = self.redact_text(p.text)
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text:
                                    p.text = self.redact_text(p.text)

        # process body para
        for p in doc.paragraphs:
            if p.text:
                p.text = self.redact_text(p.text)

        # process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text:
                            p.text = self.redact_text(p.text)

        # remove images
        images_removed = self.remove_image_pii(doc)
        print(f"Replaced {images_removed} embedded images with blank placeholder.")

        # ensure output dir
        out_dir = os.path.dirname(output_path)
        if out_dir:
            os.makedirs(out_dir, exist_ok=True)

        doc.save(output_path)
        print(f"Sanitized document successfully saved to: {output_path}")
        return True

    def redact_docx(self, input_path: str, output_path: str):
        return self.process_docx(input_path, output_path)

    def detect_and_replace_text(self, text: str) -> str:
        return self.redact_text(text)


ProductionPIIRedactor = DocumentPIIRedactor
PIIRedactor = DocumentPIIRedactor
RobustPIIRedactor = DocumentPIIRedactor


if __name__ == "__main__":
    default_input = os.path.join("input", "Red Herring Prospectus.docx")
    if not os.path.exists(default_input) and os.path.exists("Red Herring Prospectus.docx"):
        default_input = "Red Herring Prospectus.docx"

    input_file = sys.argv[1] if len(sys.argv) > 1 else default_input
    output_file = sys.argv[2] if len(sys.argv) > 2 else os.path.join("output", "redacted_output.docx")

    print(f"Input Document : {input_file}")
    print(f"Output Target  : {output_file}")

    redactor = DocumentPIIRedactor()
    redactor.process_docx(input_file, output_file)